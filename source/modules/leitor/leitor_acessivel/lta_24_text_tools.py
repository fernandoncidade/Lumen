from PySide6.QtWidgets import QFileDialog, QMessageBox
from PySide6.QtCore import QCoreApplication
from PySide6.QtGui import QTextCursor, QTextBlockFormat, QTextListFormat
from source.utils.LogManager import LogManager
import os

logger = LogManager.get_logger()

def criar_texto(self, confirmar=True):
    try:
        try:
            if getattr(self, "_content_stack", None):
                self._content_stack.setCurrentIndex(0)

        except Exception as e:
            logger.debug(f"Erro ao resetar content stack: {e}", exc_info=True)

        try:
            cur_text = self.texto_area.toPlainText()
            if confirmar and cur_text and cur_text.strip():
                from PySide6.QtWidgets import QMessageBox
                resp = QMessageBox.question(
                    self,
                    QCoreApplication.translate("App", "Criar Texto"),
                    QCoreApplication.translate("App", "Deseja criar um novo texto? O conteúdo atual será descartado."),
                    QMessageBox.Yes | QMessageBox.No
                )
                if resp != QMessageBox.Yes:
                    return

        except Exception as e:
            logger.debug(f"Erro ao confirmar criação de novo texto: {e}", exc_info=True)

        self.texto_area.clear()
        try:
            from source.utils.FontManager import FontManager
            self.texto_area.document().setDefaultFont(FontManager.get_font())

        except Exception as e:
            logger.debug(f"Erro ao definir fonte padrão: {e}", exc_info=True)

        self.texto_area.setFocus()
        logger.debug("Novo texto criado no Leitor Acessível")

    except Exception as e:
        logger.error(f"Erro ao criar texto: {e}", exc_info=True)
        try:
            QMessageBox.critical(self, QCoreApplication.translate("App", "Erro"), str(e))

        except Exception as e:
            logger.debug(f"Erro ao mostrar mensagem de erro: {e}", exc_info=True)

def salvar_como(self):
    try:
        filtro = (QCoreApplication.translate("App", "Arquivos de texto (*.txt);;Arquivos PDF (*.pdf);;Documento do Word (*.docx)"))
        path, selfilter = QFileDialog.getSaveFileName(self, QCoreApplication.translate("App", "Salvar Como"), "", filtro)
        if not path:
            return

        text = self.texto_area.toPlainText()
        _, ext = os.path.splitext(path)
        ext = ext.lower()

        if ext == ".txt":
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)

            QMessageBox.information(self, QCoreApplication.translate("App", "Salvo"), QCoreApplication.translate("App", "Arquivo TXT salvo com sucesso."))

        elif ext == ".docx":
            try:
                from docx import Document
                doc = Document()
                for p in text.split("\n\n"):
                    doc.add_paragraph(p.replace("\n", " "))

                doc.save(path)
                QMessageBox.information(self, QCoreApplication.translate("App", "Salvo"), QCoreApplication.translate("App", "Arquivo DOCX salvo com sucesso."))

            except Exception as e:
                logger.error(f"Erro ao salvar DOCX: {e}", exc_info=True)
                QMessageBox.warning(self, QCoreApplication.translate("App", "Aviso"), QCoreApplication.translate("App", "Não foi possível salvar DOCX. Verifique se 'python-docx' está instalado."))

        elif ext == ".pdf":
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas
                c = canvas.Canvas(path, pagesize=A4)
                width, height = A4
                margin = 40
                y = height - margin
                lines = text.splitlines()
                for line in lines:
                    if not line:
                        y -= 12
                        if y < margin:
                            c.showPage()
                            y = height - margin

                        continue

                    chunks = []
                    while line:
                        maxchars = 95
                        chunk = line[:maxchars]
                        if len(line) > maxchars:
                            last_space = chunk.rfind(" ")
                            if last_space > 10:
                                chunk = line[:last_space]
                                line = line[last_space+1:]

                            else:
                                line = line[maxchars:]

                        else:
                            line = ""

                        chunks.append(chunk)

                    for ch in chunks:
                        c.drawString(margin, y, ch)
                        y -= 12
                        if y < margin:
                            c.showPage()
                            y = height - margin

                c.save()
                QMessageBox.information(self, QCoreApplication.translate("App", "Salvo"), QCoreApplication.translate("App", "Arquivo PDF salvo com sucesso."))

            except Exception as e:
                logger.debug(f"reportlab indisponível ou falha ao gerar PDF: {e}", exc_info=True)
                try:
                    tmp_docx = os.path.splitext(path)[0] + "_tmp_docx_for_pdf.docx"
                    from docx import Document
                    doc = Document()
                    for p in text.split("\n\n"):
                        doc.add_paragraph(p.replace("\n", " "))

                    doc.save(tmp_docx)
                    try:
                        import docx2pdf
                        docx2pdf.convert(tmp_docx, path)
                        os.remove(tmp_docx)
                        QMessageBox.information(self, QCoreApplication.translate("App", "Salvo"), QCoreApplication.translate("App", "Arquivo PDF salvo com sucesso."))

                    except Exception as e2:
                        logger.debug(f"docx2pdf falhou: {e2}", exc_info=True)
                        if os.path.exists(tmp_docx):
                            os.remove(tmp_docx)

                        raise e

                except Exception as e3:
                    logger.error(f"Erro ao gerar PDF: {e3}", exc_info=True)
                    QMessageBox.warning(self, QCoreApplication.translate("App", "Erro"), QCoreApplication.translate("App", "Não foi possível gerar PDF. Instale 'reportlab' ou 'docx2pdf'."))

        else:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)

                QMessageBox.information(self, QCoreApplication.translate("App", "Salvo"), QCoreApplication.translate("App", "Arquivo salvo (tratado como TXT)."))

            except Exception as e:
                logger.error(f"Erro ao salvar arquivo sem extensão: {e}", exc_info=True)
                QMessageBox.warning(self, QCoreApplication.translate("App", "Erro"), QCoreApplication.translate("App", "Não foi possível salvar o arquivo."))

    except Exception as e:
        logger.error(f"Erro em salvar_como: {e}", exc_info=True)
        try:
            QMessageBox.critical(self, QCoreApplication.translate("App", "Erro"), str(e))

        except Exception as e:
            logger.debug(f"Erro ao mostrar mensagem de erro: {e}", exc_info=True)

def toggle_bullets(self):
    try:
        cursor = self.texto_area.textCursor()
        if not cursor:
            return

        fmt = QTextListFormat()
        fmt.setStyle(QTextListFormat.ListDisc)
        cursor.createList(fmt)

    except Exception as e:
        logger.error(f"Erro ao alternar marcadores: {e}", exc_info=True)

def set_line_spacing(self, spacing: float):
    try:
        if spacing <= 0:
            return

        doc = self.texto_area.document()

        block_fmt = QTextBlockFormat()
        height_value = float(spacing * 100.0)
        height_type = 1
        block_fmt.setLineHeight(height_value, height_type)

        block = doc.firstBlock()
        while block.isValid():
            try:
                bc = QTextCursor(block)
                bc.setBlockFormat(block_fmt)

            except Exception as e:
                logger.debug(f"Falha ao aplicar formato em um bloco: {e}", exc_info=True)

            block = block.next()

        logger.debug(f"Espaçamento de linhas aplicado: {spacing}")

    except Exception as e:
        logger.error(f"Erro ao definir espaçamento: {e}", exc_info=True)

def set_indent(self, indent: int):
    try:
        doc = self.texto_area.document()
        cursor = QTextCursor(doc)
        cursor.select(QTextCursor.Document)
        block_fmt = QTextBlockFormat()
        block_fmt.setLeftMargin(indent)
        cursor.mergeBlockFormat(block_fmt)
        logger.debug(f"Recuo aplicado: {indent}px")

    except Exception as e:
        logger.error(f"Erro ao definir recuo: {e}", exc_info=True)

def set_margins(self, margin: float):
    try:
        doc = self.texto_area.document()
        doc.setDocumentMargin(margin)
        logger.debug(f"Margem do documento definida: {margin}px")

    except Exception as e:
        logger.error(f"Erro ao definir margens: {e}", exc_info=True)
